#!/usr/bin/env python3
"""Generate Word document for LLC formation brief."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Title
title = doc.add_heading('Stryder Labs LLC - Formation Brief', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle/meta
meta = doc.add_paragraph()
meta.add_run('Information Package for Legal Counsel').italic = True
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
meta2 = doc.add_paragraph()
meta2.add_run('Prepared by: ').bold = True
meta2.add_run('Abraham Pena (Founder)')
meta2 = doc.add_paragraph()
meta2.add_run('Date: ').bold = True
meta2.add_run('December 10, 2025')

doc.add_paragraph()
doc.add_paragraph('─' * 50)

# Section 1: Company Overview
doc.add_heading('1. Company Overview', level=1)

doc.add_heading('Basic Information', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
data = [
    ('Proposed Entity Name', 'Stryder Labs LLC'),
    ('Entity Type', 'Limited Liability Company (LLC)'),
    ('State of Formation', 'TO BE DETERMINED - seeking counsel on Delaware vs. Florida'),
    ('Principal Place of Business', '7225 Crossroads Garden Drive Apt 4216, Orlando, FL 32821'),
    ('Organizer/Sole Member', 'Abraham Pena'),
]
for i, (field, value) in enumerate(data):
    table.rows[i].cells[0].text = field
    table.rows[i].cells[1].text = value
    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_heading('Business Purpose', level=2)
doc.add_paragraph('Stryder Labs LLC is a technology company that:')
doc.add_paragraph('1. Develops and operates the LEAG Platform - a multi-tenant SaaS platform for youth sports organizations', style='List Number')
doc.add_paragraph('2. Provides software licensing, platform hosting, and related technology services', style='List Number')
doc.add_paragraph('3. Conducts all business activities of the sole member related to software development and technology services', style='List Number')

doc.add_heading('Primary Product: LEAG Platform', level=2)
doc.add_paragraph('The LEAG Platform is a white-label SaaS solution providing:')
items = [
    'Team/organization management',
    'E-commerce (merchandise, event tickets)',
    'Member portals with player management',
    'Payment processing (Stripe integration)',
    'Content management (Wagtail CMS)',
    'Event registration and scheduling',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Section 2: NJ Stars Relationship
doc.add_heading('2. Initial Client Relationship: NJ Stars Elite AAU', level=1)

doc.add_paragraph('NJ Stars Elite AAU Basketball is the first tenant/client using the LEAG Platform. This relationship serves as the proof-of-concept for the platform\'s multi-tenant capabilities.')

doc.add_heading('Key Party Information', level=2)
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
data = [
    ('Client Entity', 'NJ Stars Elite AAU (or Kenneth Andrade\'s registered business entity)'),
    ('Client Contact', 'Kenneth Andrade (Owner)'),
    ('Relationship Type', 'Platform licensing + revenue sharing'),
]
for i, (field, value) in enumerate(data):
    table.rows[i].cells[0].text = field
    table.rows[i].cells[1].text = value
    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_heading('Revenue Structure - Phase 1: Platform Fee Model (Current)', level=2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
headers = ['Revenue Stream', 'Platform Fee (Stryder Labs)', 'Client Retention (NJ Stars)']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
data = [
    ('Events & Merchandise Sales', '20%', '80%'),
    ('Coach Services (Private Lessons)', '5%', '95%'),
]
for i, (stream, platform, client) in enumerate(data, 1):
    table.rows[i].cells[0].text = stream
    table.rows[i].cells[1].text = platform
    table.rows[i].cells[2].text = client

doc.add_paragraph()
doc.add_heading('Phase 2: Subscription/Usage Model (Future Transition)', level=2)
doc.add_paragraph('The platform fee model will transition to a subscription or usage-based model when:')
doc.add_paragraph('The LEAG Platform acquires additional tenants, OR', style='List Bullet')
doc.add_paragraph('Platform revenue from other sources is sufficient to cover development/operational costs', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Proof of Concept Phase: ').bold = True
p.add_run('6-12 months with NJ Stars as the sole tenant')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Transition Triggers:').bold = True
doc.add_paragraph('To be evaluated based on revenue trajectory after proof of concept phase (6-12 months)', style='List Bullet')
doc.add_paragraph('Transition timing to be mutually agreed upon by both parties', style='List Bullet')
doc.add_paragraph('No fixed MRR or tenant count thresholds predetermined', style='List Bullet')
doc.add_paragraph('Periodic reviews (quarterly recommended) to assess platform sustainability', style='List Bullet')

# Section 3: Revenue Interest (LEAG Platform Only)
doc.add_heading('3. Revenue Interest: Kenneth Andrade / NJ Stars', level=1)

doc.add_paragraph('As recognition of NJ Stars\' role as the founding client and development partner, the following consideration is proposed:')

# Important callout
p = doc.add_paragraph()
p.add_run('IMPORTANT: ').bold = True
p.add_run('This revenue interest is specifically tied to the LEAG Platform product only, NOT to Stryder Labs LLC as a whole. Stryder Labs may pursue other ventures, products, or services that would not be subject to this arrangement.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Trigger Event: ').bold = True
p.add_run('Upon completion of the Phase 1 to Phase 2 transition (platform fee phase-out)')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Revenue/Profit Interest Grant:').bold = True
doc.add_paragraph('Recipient: Kenneth Andrade personally, OR the NJ Stars registered business entity (to be determined)', style='List Bullet')
doc.add_paragraph('Amount: 3-5% of LEAG Platform net revenue/profits (exact percentage and structure to be finalized)', style='List Bullet')
doc.add_paragraph('Scope: Limited exclusively to revenue generated by the LEAG Platform', style='List Bullet')
doc.add_paragraph('Type: Profit interest or revenue royalty (NOT equity in Stryder Labs LLC itself)', style='List Bullet')
doc.add_paragraph('Duration: Perpetual OR time-limited (to be determined)', style='List Bullet')

doc.add_heading('Conditions for Revenue Interest Grant', level=2)
doc.add_paragraph('The revenue interest grant is contingent upon:')
doc.add_paragraph('NJ Stars remaining an active, paying client through the transition period', style='List Number')
doc.add_paragraph('No material breach of the client agreement by NJ Stars', style='List Number')
doc.add_paragraph('Execution of appropriate membership interest purchase agreement', style='List Number')

doc.add_heading('Alternative Structures to Consider', level=2)
doc.add_paragraph('Profit Interest: Instead of direct equity, grant a profit interest that only participates in future appreciation', style='List Number')
doc.add_paragraph('Phantom Equity: Provides economic benefits of ownership without actual ownership', style='List Number')
doc.add_paragraph('Warrant/Option: Right to purchase equity at a fixed price in the future', style='List Number')
doc.add_paragraph('Revenue Royalty: Ongoing small percentage of platform revenue instead of equity', style='List Number')

# Section 4: Intellectual Property
doc.add_heading('4. Intellectual Property', level=1)

doc.add_heading('IP Ownership', level=2)
doc.add_paragraph('All intellectual property related to the LEAG Platform is owned exclusively by Stryder Labs LLC, including:')
items = [
    'Source code (frontend and backend)',
    'Design systems and UI/UX patterns',
    'Database schemas and architectures',
    'APIs and integration frameworks',
    'Documentation and technical specifications',
    'Trademarks (LEAG, Stryder Labs, platform branding)',
    'Trade secrets (algorithms, business logic)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('IP Assignment', level=2)
doc.add_paragraph('Any code, designs, or other work product created by Abraham Pena for the LEAG Platform prior to LLC formation should be formally assigned to Stryder Labs LLC.')

doc.add_heading('Client Data vs Platform IP', level=2)
doc.add_paragraph('Platform IP: Owned by Stryder Labs LLC', style='List Bullet')
doc.add_paragraph('Client Data: Each tenant owns their own data (user information, transactions, content)', style='List Bullet')
doc.add_paragraph('Aggregated/Anonymized Data: Stryder Labs may use for platform improvement', style='List Bullet')

# Section 5: Requested Legal Documents
doc.add_heading('5. Requested Legal Documents', level=1)

doc.add_heading('Formation Documents', level=2)
doc.add_paragraph('Articles of Organization - File with state', style='List Number')
doc.add_paragraph('Operating Agreement - Single-member LLC operating agreement', style='List Number')
doc.add_paragraph('EIN Application - Federal tax ID', style='List Number')

doc.add_heading('Business Agreements (Priority Order)', level=2)
p = doc.add_paragraph()
p.add_run('1. Platform Services Agreement (NJ Stars)').bold = True
doc.add_paragraph('Licensing terms', style='List Bullet')
doc.add_paragraph('Revenue sharing structure', style='List Bullet')
doc.add_paragraph('Data ownership', style='List Bullet')
doc.add_paragraph('Service level commitments', style='List Bullet')
doc.add_paragraph('Termination provisions', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('2. Equity Grant Agreement / Letter of Intent').bold = True
doc.add_paragraph('Terms for Kenneth Andrade equity consideration', style='List Bullet')
doc.add_paragraph('Trigger conditions', style='List Bullet')
doc.add_paragraph('Vesting terms (if applicable)', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('3. IP Assignment Agreement').bold = True
doc.add_paragraph('Transfer of pre-formation IP to LLC', style='List Bullet')

doc.add_heading('Future Documents (As Needed)', level=2)
items = [
    'Standard Terms of Service (for platform users)',
    'Privacy Policy',
    'Multi-tenant Platform Agreement (template)',
    'Independent Contractor Agreements',
    'Employment Agreements (when hiring)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Section 6: Questions for Legal Counsel
doc.add_heading('6. Questions for Legal Counsel', level=1)

doc.add_heading('Formation Questions', level=2)
doc.add_paragraph('What state should we form in? (Delaware vs. home state considerations)', style='List Number')
doc.add_paragraph('Should we elect S-Corp tax treatment for potential tax savings?', style='List Number')
doc.add_paragraph('What registered agent services do you recommend?', style='List Number')

doc.add_heading('NJ Stars Relationship Questions', level=2)
doc.add_paragraph('What\'s the best structure for the equity grant to Kenneth Andrade?', style='List Number')
doc.add_paragraph('How should we document the platform fee to subscription transition?', style='List Number')
doc.add_paragraph('Should the equity grant be to Kenneth personally or his business entity?', style='List Number')
doc.add_paragraph('What protections should we include for both parties?', style='List Number')

doc.add_heading('IP Questions', level=2)
doc.add_paragraph('Do we need formal IP assignment documents for pre-formation work?', style='List Number')
doc.add_paragraph('How do we protect the platform IP while licensing to tenants?', style='List Number')
doc.add_paragraph('What trademark filings should we consider?', style='List Number')

doc.add_heading('Revenue/Tax Questions', level=2)
doc.add_paragraph('How should we structure revenue recognition for the platform fee model?', style='List Number')
doc.add_paragraph('What accounting practices should we establish from day one?', style='List Number')
doc.add_paragraph('Any state tax nexus considerations for SaaS?', style='List Number')

# Section 7: Proposed Timeline
doc.add_heading('7. Proposed Timeline', level=1)

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
headers = ['Phase', 'Target Date', 'Actions']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
data = [
    ('1. Formation', 'December 2025 (ASAP)', 'File Articles, create Operating Agreement, obtain EIN'),
    ('2. IP Assignment', '+1 week after formation', 'Execute IP assignment to LLC'),
    ('3. NJ Stars Agreement', '+2 weeks after formation', 'Execute platform services agreement'),
    ('4. Equity LOI', '+2 weeks after formation', 'Document equity consideration terms'),
    ('5. Bank Account', '+3 weeks after formation', 'Open business banking'),
    ('6. Accounting Setup', '+4 weeks after formation', 'Establish bookkeeping practices'),
]
for i, (phase, date, actions) in enumerate(data, 1):
    table.rows[i].cells[0].text = phase
    table.rows[i].cells[1].text = date
    table.rows[i].cells[2].text = actions

# Section 8: Financial Projections
doc.add_heading('8. Financial Projections', level=1)

doc.add_paragraph().add_run('Platform has not yet launched - no historical revenue data available.').italic = True

doc.add_paragraph()
doc.add_paragraph('Revenue projections will be developed post-launch based on:')
doc.add_paragraph('Event registration volume and pricing', style='List Bullet')
doc.add_paragraph('Merchandise sales velocity', style='List Bullet')
doc.add_paragraph('Coach services adoption (private lessons)', style='List Bullet')

doc.add_heading('Growth Projections', level=2)
doc.add_paragraph('Months 1-6 (or up to 12): NJ Stars only (proof of concept phase)', style='List Bullet')
doc.add_paragraph('Post-POC Year 1: 2-3 additional tenants, evaluate transition triggers', style='List Bullet')
doc.add_paragraph('Post-POC Year 2: 5-10 tenants, potential subscription model transition', style='List Bullet')

# Section 9: Contact Information
doc.add_heading('9. Contact Information', level=1)

doc.add_heading('Founder', level=2)
p = doc.add_paragraph()
p.add_run('Abraham Pena').bold = True
doc.add_paragraph('Email: pena.abe@gmail.com')
doc.add_paragraph('Phone: (949) 836-6285')
doc.add_paragraph('Address: 7225 Crossroads Garden Drive Apt 4216, Orlando, FL 32821')

doc.add_heading('Client Contact (NJ Stars)', level=2)
p = doc.add_paragraph()
p.add_run('Kenneth Andrade').bold = True
doc.add_paragraph('Email: kennethandrade89@gmail.com')
doc.add_paragraph('Phone: (201) 468-1445')
doc.add_paragraph('Organization: NJ Stars Elite AAU')

# Section 10: Attachments
doc.add_heading('10. Attachments / Additional Materials', level=1)

doc.add_paragraph('Please let me know if you need any of the following:')
doc.add_paragraph('☐ Copy of government-issued ID')
doc.add_paragraph('☐ Proof of address')
doc.add_paragraph('☐ NJ Stars business registration (if available)')
doc.add_paragraph('☐ Platform technical overview')
doc.add_paragraph('☐ Any existing written agreements with Kenneth Andrade')
doc.add_paragraph('☐ Desired operating agreement provisions')

# Footer
doc.add_paragraph()
doc.add_paragraph('─' * 50)
p = doc.add_paragraph()
p.add_run('Confidentiality: ').bold = True
p.add_run('This document contains business strategy and should be treated as confidential.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Document prepared December 10, 2025').italic = True

# Save
doc.save('/Users/abe/Desktop/Projects/leag/documents/Stryder_Labs_LLC_Formation_Brief.docx')
print('Document created successfully!')
